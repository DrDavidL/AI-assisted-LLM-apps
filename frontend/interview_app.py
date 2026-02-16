"""Streamlit app for voice-based patient interview practice."""

from __future__ import annotations

import base64
import io
import os
from datetime import datetime, timezone
from pathlib import Path
import random
import uuid

from docx import Document
from dotenv import load_dotenv
import openai
import streamlit as st

from api_client import get_case_by_number, save_transcript

load_dotenv(Path(__file__).resolve().parent.parent / ".env")

st.set_page_config(page_title="Patient Interview", layout="wide")

OPENAI_AUDIO_MODEL = os.getenv("OPENAI_AUDIO_MODEL", "gpt-4o-audio-preview")
VOICES = ["alloy", "ash", "ballad", "coral", "echo", "fable", "onyx", "nova", "sage", "shimmer"]


# ── Helpers ──────────────────────────────────────────────────────────────────


def _build_patient_system_prompt(case: dict) -> str:
    """Build a system prompt that instructs the AI to role-play as the patient."""
    demo = case.get("demographics") or {}
    hpi = case.get("chief_complaint_hpi") or {}
    social = case.get("social_history") or {}
    meds = case.get("medications") or []
    allergies = case.get("allergies") or []
    vitals = case.get("vitals") or {}
    pmh = case.get("past_medical_history") or {}

    age = demo.get("age", "unknown")
    sex = demo.get("sex", "unknown")
    language = demo.get("preferred_language", "English")

    med_list = ", ".join(
        f"{m.get('name', '?')} {m.get('dose', '')}".strip() for m in meds
    ) or "none"
    allergy_list = ", ".join(
        f"{a.get('substance', '?')} ({a.get('reaction', 'unknown reaction')})" for a in allergies
    ) or "none"
    conditions = ", ".join(pmh.get("conditions", [])) or "none"

    return f"""You are a patient in a medical interview. Stay in character at all times.

PERSONA:
- Age: {age}, Sex: {sex}, Preferred language: {language}
- You speak like a regular patient, not a clinician. Use everyday language.

YOUR SYMPTOMS (what you know and can share):
- Chief complaint: {hpi.get('chief_complaint', 'not specified')}
- Story: {hpi.get('hpi_narrative', 'not specified')}
- Onset: {hpi.get('onset', 'not sure')}
- Duration: {hpi.get('duration', 'not sure')}
- Severity: {hpi.get('severity', 'not sure')}
- What makes it worse: {', '.join(hpi.get('aggravating_factors', [])) or 'nothing specific'}
- What makes it better: {', '.join(hpi.get('alleviating_factors', [])) or 'nothing specific'}
- Other symptoms: {', '.join(hpi.get('associated_symptoms', [])) or 'none'}

YOUR BACKGROUND (reveal when asked):
- Medical conditions: {conditions}
- Current medications: {med_list}
- Allergies: {allergy_list}
- Tobacco: {social.get('tobacco', 'not specified')}
- Alcohol: {social.get('alcohol', 'not specified')}
- Drugs: {social.get('drugs', 'not specified')}
- Occupation: {social.get('occupation', 'not specified')}
- Living situation: {social.get('living_situation', 'not specified')}

YOUR VITALS (you may mention how you feel, but do NOT cite exact numbers):
- Pain level: {vitals.get('pain_scale', 'not specified')}/10

RULES:
- You do NOT know your diagnosis. You only know your symptoms and history.
- Answer questions naturally and conversationally as a real patient would.
- If the student asks something you don't know, say so naturally.
- Keep answers concise — 1-3 sentences unless the question requires more detail.
- Do NOT volunteer all information at once; let the student ask."""


def _transcribe_audio(audio_bytes: bytes) -> str:
    """Transcribe user audio via OpenAI Whisper. Returns text or fallback string."""
    client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "question.wav"
    result = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    return result.text.strip() or "(inaudible)"


def _get_ai_response(messages: list, voice: str) -> tuple[str, bytes, str]:
    """Call OpenAI Chat Completions with audio modalities. Returns (transcript, wav_bytes, audio_id)."""
    client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    completion = client.chat.completions.create(
        model=OPENAI_AUDIO_MODEL,
        modalities=["text", "audio"],
        audio={"voice": voice, "format": "wav"},
        messages=messages,
    )
    choice = completion.choices[0].message
    transcript = choice.audio.transcript
    wav_bytes = base64.b64decode(choice.audio.data)
    audio_id = choice.audio.id
    return transcript, wav_bytes, audio_id


def _build_transcript_docx(case: dict, history: list[dict]) -> bytes:
    """Build a Word document from the interview history and return raw bytes."""
    doc = Document()
    doc.add_heading(case.get("case_title", "Untitled Case"), level=1)
    doc.add_paragraph(
        f"Case #{case.get('case_number', '—')} | "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    doc.add_paragraph("")

    for entry in history:
        label = "Student" if entry["role"] == "user" else "Patient"
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(entry["text"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Session state init ───────────────────────────────────────────────────────

if "interview_case" not in st.session_state:
    st.session_state["interview_case"] = None
if "interview_messages" not in st.session_state:
    st.session_state["interview_messages"] = []
if "interview_voice" not in st.session_state:
    st.session_state["interview_voice"] = None
if "interview_history" not in st.session_state:
    st.session_state["interview_history"] = []  # list of {"role", "text", "audio_bytes"|None}
if "interview_conversation_id" not in st.session_state:
    st.session_state["interview_conversation_id"] = None

# ── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.header("Patient Interview")

    case_num = st.number_input("Case number", min_value=1, step=1, key="iv_case_num")

    if st.button("Load Case", type="primary"):
        try:
            case = get_case_by_number(case_num)
            voice = random.choice(VOICES)
            st.session_state["interview_case"] = case
            st.session_state["interview_voice"] = voice
            st.session_state["interview_conversation_id"] = str(uuid.uuid4())
            st.session_state["interview_messages"] = [
                {"role": "system", "content": _build_patient_system_prompt(case)},
            ]
            st.session_state["interview_history"] = []
            st.success(f"Loaded: {case.get('case_title', 'Untitled')}")
            st.rerun()
        except Exception as e:
            st.error(f"Failed to load case #{case_num}: {e}")

    if st.session_state["interview_voice"]:
        st.caption(f"Voice: {st.session_state['interview_voice']}")

    if st.session_state["interview_history"]:
        docx_bytes = _build_transcript_docx(
            st.session_state["interview_case"],
            st.session_state["interview_history"],
        )
        st.download_button(
            label="Download Transcript (.docx)",
            data=docx_bytes,
            file_name=f"transcript_case_{st.session_state['interview_case'].get('case_number', 'unknown')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if st.session_state["interview_case"] and st.button("Reset Interview"):
        st.session_state["interview_case"] = None
        st.session_state["interview_messages"] = []
        st.session_state["interview_voice"] = None
        st.session_state["interview_history"] = []
        st.session_state["interview_conversation_id"] = None
        st.rerun()

# ── Main area ────────────────────────────────────────────────────────────────

case = st.session_state["interview_case"]

if not case:
    st.title("Patient Interview Practice")
    st.info("Enter a case number in the sidebar and click **Load Case** to begin.")
    st.stop()

# Case summary header
st.title(case.get("case_title", "Untitled Case"))
demo = case.get("demographics") or {}
st.caption(
    f"Case #{case.get('case_number', '—')} | "
    f"{demo.get('age', '?')}yo {demo.get('sex', '?')} | "
    f"Specialty: {case.get('specialty', '')} | "
    f"Difficulty: {case.get('difficulty', '')}"
)

st.divider()

# Conversation history
for entry in st.session_state["interview_history"]:
    role_label = "You" if entry["role"] == "user" else "Patient"
    with st.chat_message("user" if entry["role"] == "user" else "assistant"):
        st.markdown(f"**{role_label}:** {entry['text']}")
        if entry.get("audio_bytes"):
            st.audio(entry["audio_bytes"], format="audio/wav")

# Audio input
audio_input = st.audio_input("Record your question")

if audio_input is not None:
    # Prevent re-processing the same audio on rerun by tracking file id
    audio_id_key = f"{audio_input.file_id}"
    if st.session_state.get("_last_audio_id") == audio_id_key:
        st.stop()
    st.session_state["_last_audio_id"] = audio_id_key

    audio_bytes = audio_input.read()
    audio_b64 = base64.b64encode(audio_bytes).decode("utf-8")

    # Transcribe the student's audio via Whisper
    with st.spinner("Transcribing your question..."):
        try:
            user_text = _transcribe_audio(audio_bytes)
        except Exception:
            user_text = "(audio question)"

    # Add user audio message to OpenAI conversation
    st.session_state["interview_messages"].append({
        "role": "user",
        "content": [
            {
                "type": "input_audio",
                "input_audio": {"data": audio_b64, "format": "wav"},
            }
        ],
    })

    # Add user turn with transcribed text to display history
    st.session_state["interview_history"].append({
        "role": "user",
        "text": user_text,
        "audio_bytes": audio_bytes,
    })

    with st.spinner("Patient is responding..."):
        try:
            transcript, wav_bytes, resp_audio_id = _get_ai_response(
                st.session_state["interview_messages"],
                st.session_state["interview_voice"],
            )

            # Append assistant message using audio id reference for multi-turn
            st.session_state["interview_messages"].append({
                "role": "assistant",
                "audio": {"id": resp_audio_id},
            })

            # Add patient response to display history
            st.session_state["interview_history"].append({
                "role": "assistant",
                "text": transcript,
                "audio_bytes": wav_bytes,
            })

            # Auto-save transcript to DB
            try:
                save_transcript(
                    conversation_id=st.session_state["interview_conversation_id"],
                    case_number=case.get("case_number", 0),
                    transcript=[
                        {"role": e["role"], "text": e["text"]}
                        for e in st.session_state["interview_history"]
                    ],
                )
            except Exception:
                pass  # best-effort save; don't block the interview

            st.rerun()
        except Exception as e:
            st.error(f"Error getting response: {e}")
            # Remove the failed user and history entries so conversation stays consistent
            st.session_state["interview_messages"].pop()
            st.session_state["interview_history"].pop()
